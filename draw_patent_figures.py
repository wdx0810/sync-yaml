# -*- coding: utf-8 -*-
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import os

OUT = r"C:\AI\project-configmap\project-configmap\patent_figures"
os.makedirs(OUT, exist_ok=True)
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False

def box(ax, x, y, w, h, text, fc='#dbeafe', ec='#2563eb', fs=8.5):
    ax.add_patch(FancyBboxPatch((x-w/2, y-h/2), w, h,
        boxstyle="round,pad=0.05", facecolor=fc, edgecolor=ec))
    ax.text(x, y, text, ha='center', va='center', fontsize=fs, wrap=True)

def arrow(ax, x1, y1, x2, y2):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle='->', color='#374151', lw=1.3))

def save(fig, name):
    fig.savefig(os.path.join(OUT, name), dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  {name}")

# Fig 1
print("Fig 1...")
fig, ax = plt.subplots(figsize=(11, 7))
ax.set_xlim(0,11); ax.set_ylim(0,7); ax.axis('off')
ax.set_title('图1 双向同步系统整体架构图', fontsize=13, fontweight='bold')
ax.add_patch(FancyBboxPatch((0.3,1.8), 10.4, 4.8, boxstyle="round,pad=0.1", fc='#f8fafc', ec='#2563eb', lw=2))
ax.text(5.5, 6.3, 'YAML Sync 双向同步系统', ha='center', fontsize=12, fontweight='bold', color='#1e40af')
for label, cx in [('Web前端', 1.5), ('API Server', 3.8), ('Task Manager', 6.2), ('History', 8.8)]:
    box(ax, cx, 5.5, 1.8, 0.7, label, '#dbeafe', '#3b82f6')
box(ax, 5.5, 4.2, 8, 0.9, 'Generic Syncer: YAML Parser + GVR解析 + 三级清理器 + JSON归一化比较', '#fef3c7', '#d97706', 9)
box(ax, 2.5, 2.8, 2.5, 0.8, 'GitLab Client\n(原子提交)', '#dcfce7', '#16a34a')
box(ax, 8.5, 2.8, 2.5, 0.8, 'K8s Dynamic Client\n(SSA)', '#dcfce7', '#16a34a')
box(ax, 2.5, 0.8, 2.5, 0.8, 'GitLab 仓库', '#fce7f3', '#db2777')
box(ax, 8.5, 0.8, 2.5, 0.8, 'K8s 集群', '#fce7f3', '#db2777')
arrow(ax, 5.5, 5.1, 5.5, 4.65)
arrow(ax, 3.5, 3.75, 2.5, 3.2)
arrow(ax, 7.5, 3.75, 8.5, 3.2)
arrow(ax, 2.5, 2.4, 2.5, 1.2)
arrow(ax, 8.5, 2.4, 8.5, 1.2)
ax.annotate('', xy=(4, 0.8), xytext=(7, 0.8), arrowprops=dict(arrowstyle='<->', color='#059669', lw=2.5))
ax.text(5.5, 0.8, '双向', ha='center', fontsize=10, color='#059669', fontweight='bold')
save(fig, 'fig1.png')

# Fig 2
print("Fig 2...")
fig, ax = plt.subplots(figsize=(7, 11))
ax.set_xlim(0,7); ax.set_ylim(0,11); ax.axis('off')
ax.set_title('图2 正向同步方法流程图', fontsize=13, fontweight='bold')
steps = ['开始', '从GitLab获取YAML文件', 'sigs.k8s.io/yaml解析', 'GVR解析器确定资源类型',
         '对两侧执行三级清理', 'JSON归一化比较', '生成变更预览', '用户审核(逐项勾选)',
         'Server-Side Apply', '结束']
colors = ['#e5e7eb','#dbeafe','#dbeafe','#dbeafe','#fef3c7','#fef3c7','#dcfce7','#dcfce7','#dcfce7','#e5e7eb']
for i, (s, c) in enumerate(zip(steps, colors)):
    y = 10 - i * 1.0
    box(ax, 3.5, y, 3.2, 0.7, s, c, '#374151')
    if i < len(steps)-1:
        arrow(ax, 3.5, y-0.35, 3.5, y-0.65)
ax.text(6, 5, 'same\n->skip', fontsize=8, color='red')
save(fig, 'fig2.png')

# Fig 3
print("Fig 3...")
fig, ax = plt.subplots(figsize=(7, 9))
ax.set_xlim(0,7); ax.set_ylim(0,9); ax.axis('off')
ax.set_title('图3 反向同步方法流程图', fontsize=13, fontweight='bold')
steps = ['开始', '确定资源类型+命名空间', 'Dynamic Client列出资源', '自动生成资源过滤',
         '三级运行时字段清理', '序列化为精简YAML', '与GitLab比较', '多文件原子提交', '结束']
colors = ['#e5e7eb','#dbeafe','#dbeafe','#fef3c7','#fef3c7','#fef3c7','#dcfce7','#dcfce7','#e5e7eb']
for i, (s, c) in enumerate(zip(steps, colors)):
    y = 8 - i * 0.9
    box(ax, 3.5, y, 3.2, 0.65, s, c, '#374151')
    if i < len(steps)-1:
        arrow(ax, 3.5, y-0.33, 3.5, y-0.57)
save(fig, 'fig3.png')

# Fig 4
print("Fig 4...")
fig, ax = plt.subplots(figsize=(9, 6))
ax.set_xlim(0,9); ax.set_ylim(0,6); ax.axis('off')
ax.set_title('图4 三级运行时字段智能清理方法', fontsize=13, fontweight='bold')
box(ax, 4.5, 5.3, 3, 0.5, '输入: K8s资源对象', '#e5e7eb', '#374151')
box(ax, 4.5, 4.2, 7, 0.8, '第一级(通用): managedFields/resourceVersion/uid/creationTimestamp/ownerReferences', '#dbeafe', '#2563eb', 8)
box(ax, 4.5, 3.0, 7, 0.8, '第二级(类型): Service:clusterIP | Deploy:progressDeadline/revisionHistoryLimit | SA:secrets', '#fef3c7', '#d97706', 8)
box(ax, 4.5, 1.8, 7, 0.8, '第三级(容器): dnsPolicy/restartPolicy/imagePullPolicy/探针默认值/kube-api-access卷', '#dcfce7', '#16a34a', 8)
box(ax, 4.5, 0.7, 3, 0.5, '输出: 可移植精简对象', '#e5e7eb', '#374151')
for y1, y2 in [(5.05,4.6),(3.8,3.4),(2.6,2.2),(1.4,0.95)]:
    arrow(ax, 4.5, y1, 4.5, y2)
ax.text(8.2, 0.7, '关键:\nint/int64/float64\n类型感知', fontsize=8, color='red',
    bbox=dict(fc='#fef2f2', ec='red', boxstyle='round'))
save(fig, 'fig4.png')

# Fig 5
print("Fig 5...")
fig, ax = plt.subplots(figsize=(7, 5))
ax.set_xlim(0,7); ax.set_ylim(0,5); ax.axis('off')
ax.set_title('图5 JSON归一化比较方法', fontsize=13, fontweight='bold')
box(ax, 2, 4.2, 2.2, 0.6, 'K8s对象\n(int64)', '#dbeafe', '#2563eb')
box(ax, 5, 4.2, 2.2, 0.6, 'GitLab对象\n(float64)', '#fef3c7', '#d97706')
box(ax, 3.5, 3, 4, 0.7, '递归归一化: 全部数值->float64\n过滤nil/空map/空slice', '#f3e8ff', '#7c3aed')
box(ax, 3.5, 1.8, 3.5, 0.6, 'JSON序列化 -> 字符串比较', '#dcfce7', '#16a34a')
box(ax, 2, 0.7, 1.5, 0.5, '相同:跳过', '#dcfce7', '#16a34a')
box(ax, 5, 0.7, 1.5, 0.5, '不同:变更', '#fef2f2', '#dc2626')
arrow(ax, 2, 3.9, 3, 3.35); arrow(ax, 5, 3.9, 4, 3.35)
arrow(ax, 3.5, 2.65, 3.5, 2.1)
arrow(ax, 2.8, 1.5, 2, 0.95); arrow(ax, 4.2, 1.5, 5, 0.95)
save(fig, 'fig5.png')

# Fig 6
print("Fig 6...")
fig, ax = plt.subplots(figsize=(7, 7))
ax.set_xlim(0,7); ax.set_ylim(0,7); ax.axis('off')
ax.set_title('图6 自动生成资源过滤规则', fontsize=13, fontweight='bold')
box(ax, 3.5, 6.2, 3, 0.5, '输入: 集群资源', '#e5e7eb', '#374151')
box(ax, 3.5, 5.2, 3, 0.6, '有OwnerReferences?', '#fef3c7', '#d97706')
box(ax, 3.5, 4.0, 3, 0.6, '纯运行时类型?\n(Pod/RS/EP)', '#fef3c7', '#d97706')
box(ax, 3.5, 2.8, 3, 0.6, '系统自动创建?\n(default SA等)', '#fef3c7', '#d97706')
box(ax, 3.5, 1.5, 3, 0.6, '通过 -> 继续同步', '#dcfce7', '#16a34a')
for y1, y2 in [(5.95,5.5),(4.9,4.3),(3.7,3.1),(2.5,1.8)]:
    arrow(ax, 3.5, y1, 3.5, y2)
for y in [5.2, 4.0, 2.8]:
    ax.text(5.8, y, '是->跳过', fontsize=9, color='red')
    arrow(ax, 5.0, y, 5.5, y)
save(fig, 'fig6.png')

# Insert into docx
print("\nInserting into docx...")
from docx import Document
from docx.shared import Cm

DOCX = r"C:\AI\project-configmap\project-configmap\专利技术交底书-YAML同步系统.docx"
doc = Document(DOCX)

# Find placeholder and replace with images
for i, p in enumerate(doc.paragraphs):
    if '详细附图见附件' in p.text:
        p.text = ''
        break

# Add figures before section 5
figs = ['fig1.png','fig2.png','fig3.png','fig4.png','fig5.png','fig6.png']
captions = ['图1 双向同步系统整体架构图','图2 正向同步方法流程图','图3 反向同步方法流程图',
            '图4 三级运行时字段智能清理方法','图5 JSON归一化比较方法','图6 自动生成资源过滤规则']

for fname, cap in zip(figs, captions):
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run()
    run.add_picture(os.path.join(OUT, fname), width=Cm(14))
    c = doc.add_paragraph(cap)
    c.alignment = 1
    c.runs[0].bold = True
    doc.add_paragraph()

doc.save(DOCX)
print(f"Done! {DOCX}")
