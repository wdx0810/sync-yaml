# -*- coding: utf-8 -*-
"""Rebuild patent docx with figures in the correct position (Section 4)."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = r"C:\AI\project-configmap\project-configmap\patent_figures"
OUTPUT = r"C:\AI\project-configmap\project-configmap\专利技术交底书-YAML同步系统.docx"

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# ===== Title =====
title = doc.add_heading('专利技术交底书', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===== Info Table =====
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '发明名称'
table.rows[0].cells[1].text = '一种基于运行时字段智能清理的Kubernetes资源双向同步方法及系统'
table.rows[1].cells[0].text = '申请类型'
table.rows[1].cells[1].text = '发明'
table.rows[2].cells[0].text = '第一技术联系人'
table.rows[2].cells[1].text = ''
table.rows[3].cells[0].text = '部门'
table.rows[3].cells[1].text = ''
doc.add_paragraph()

# ===== Section 1 =====
doc.add_heading('一、技术领域', level=1)
doc.add_paragraph('本申请涉及云计算容器编排技术领域，特别涉及一种基于运行时字段智能清理的Kubernetes资源双向同步方法及系统。')

# ===== Section 2 =====
doc.add_heading('二、背景技术', level=1)
doc.add_heading('A1、与本申请相关的现有技术背景情况', level=2)
doc.add_paragraph(
    '随着云原生技术的快速发展，Kubernetes已成为容器编排的事实标准，企业普遍采用GitOps工作流将Git仓库作为基础设施配置的唯一可信源。'
    '在多集群、多环境场景下，运维团队面临两个核心需求：将Git仓库中的声明式YAML配置可靠地应用到目标集群（正向同步）；'
    '将集群中的实际运行配置导出回Git仓库用于备份、审计或跨集群迁移（反向同步）。')

doc.add_heading('A2、与本申请相关的最接近的现有技术', level=2)
doc.add_paragraph('经检索，目前市面上与Kubernetes资源同步相关的系统和专利主要包括：')
doc.add_paragraph('1. ArgoCD（Intuit公司）：最流行的GitOps工具，仅支持单向同步（Git→集群），通过对比期望状态与实际状态检测漂移。')
doc.add_paragraph('2. FluxCD（Weaveworks公司）：单向GitOps工具，集群内控制器主动从Git拉取配置。')
doc.add_paragraph('3. Google Config Sync：GKE专属配置同步组件，仅支持Google Cloud环境。')
doc.add_paragraph('4. 相关专利检索：')
doc.add_paragraph('   - CN112422555A：Kubernetes资源权限管理——关注权限，非同步')
doc.add_paragraph('   - CN119961006B：Kubernetes资源分配方法——关注调度，非配置同步')
doc.add_paragraph('   - WO2019184164A1：Kubernetes节点部署——关注部署，非YAML同步')
p = doc.add_paragraph()
run = p.add_run('经过充分检索，未发现与本申请技术方案相同或实质相同的已公开专利或论文。')
run.bold = True

doc.add_heading('A3、现有技术的缺陷和不足', level=2)
doc.add_paragraph('上述现有技术存在以下根本性缺陷：')
defects = [
    ('缺陷一：仅支持单向同步', 'ArgoCD、FluxCD均为单向工具，不支持反向导出。运维人员只能手动kubectl导出，效率低且易遗漏。'),
    ('缺陷二：往返不一致问题', 'K8s API Server自动注入大量运行时字段和默认值（resourceVersion、clusterIP、progressDeadlineSeconds=600、imagePullPolicy=IfNotPresent等50余种），导出的YAML无法直接应用到其他集群，报"field is immutable"错误。目前无工具能自动完整清理。'),
    ('缺陷三：漂移检测大量误报', 'K8s API返回int64类型数值，YAML解析器产生float64类型。reflect.DeepEqual判定int64(80)!=float64(80)，产生大量假漂移告警。ArgoCD需用户手动配置ignoreDifferences规避，复杂且易遗漏。'),
    ('缺陷四：缺乏变更审核机制', '现有工具检测到Git变更后自动应用，缺少人工审核。ArgoCD的diff视图包含大量默认值噪音，实际可用性差。'),
    ('缺陷五：自动生成资源污染Git仓库', '集群中大量控制器自动创建的资源（ReplicaSet、Pod、default SA、kube-root-ca.crt等）不应存储在Git中，现有工具缺乏过滤机制。'),
]
for title, desc in defects:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(desc)

# ===== Section 3 =====
doc.add_heading('三、发明内容', level=1)
doc.add_heading('B1、本申请所要解决的技术问题', level=2)
doc.add_paragraph('本申请解决以下技术问题：')
for prob in ['1. 单向同步无法反向导出的问题；', '2. 运行时字段导致往返不一致的问题；',
             '3. 数值类型不一致导致漂移误报的问题；', '4. 缺乏变更审核机制的问题；',
             '5. 自动生成资源污染Git仓库的问题。']:
    doc.add_paragraph(prob)

doc.add_heading('B2、本申请的技术方案及区别', level=2)
doc.add_paragraph('本申请提出一种基于运行时字段智能清理的Kubernetes资源双向同步方法，核心方案包括：')
schemes = [
    ('（一）三级运行时字段智能清理方法', '通用层删除所有资源共有的服务端字段；类型特化层按Kind删除特有默认值；容器层删除Pod模板级默认值。关键创新：数值类型感知匹配，同时支持int/int64/float64。'),
    ('（二）JSON归一化资源比较方法', '递归将所有数值统一为float64，过滤nil/空容器，JSON序列化后字符串比较，彻底消除类型不一致误报。'),
    ('（三）基于OwnerReferences的自动生成资源过滤', '通过OwnerRef检查+类型黑名单+系统资源白名单，自动过滤不应同步的资源。'),
    ('（四）预览-审核-应用三阶段正向同步', '预览阶段对两侧均执行清理后生成精确差异；用户逐项审核；通过后用原始YAML通过Server-Side Apply应用。'),
    ('（五）多文件原子提交的反向同步', '一次同步所有变更通过Git多文件API原子提交为单个commit。'),
]
for t, d in schemes:
    p = doc.add_paragraph()
    run = p.add_run(t)
    run.bold = True
    doc.add_paragraph(d)

doc.add_heading('B3、本申请的技术效果', level=2)
comp = doc.add_table(rows=8, cols=3)
comp.style = 'Table Grid'
for i, h in enumerate(['能力维度', 'ArgoCD/FluxCD', '本申请']):
    comp.rows[0].cells[i].text = h
rows = [['同步方向','仅Git→集群','Git↔集群（双向）'],['反向导出','不支持','支持'],
        ['YAML可移植性','不适用','清理后可直接应用'],['漂移检测','大量误报','零误报'],
        ['变更审核','自动应用+噪音diff','三阶段审核+精确diff'],['资源过滤','无','自动过滤'],
        ['部署复杂度','需集群内CRD','单二进制外部署']]
for i, r in enumerate(rows):
    for j, c in enumerate(r):
        comp.rows[i+1].cells[j].text = c
doc.add_paragraph()
doc.add_paragraph('量化效果：往返一致性从<10%提升到99%+；342个资源实测假变更从98个降到0个。')

doc.add_heading('B4、本申请的技术创新点', level=2)
for inn in [
    '1. 首次提出三级运行时字段智能清理架构，具有数值类型感知能力；',
    '2. 首次提出基于JSON归一化的资源实质相同性判断方法；',
    '3. 首次实现具有审核机制的Kubernetes资源双向同步系统；',
    '4. 首次提出基于OwnerReferences的复合资源过滤规则。']:
    doc.add_paragraph(inn)

# ===== Section 4: 附图 (figures go HERE) =====
doc.add_heading('四、专利附图', level=1)
doc.add_paragraph('附图说明：')
for d in ['图1为双向同步系统整体架构图','图2为正向同步方法流程图','图3为反向同步方法流程图',
          '图4为三级运行时字段智能清理方法流程图','图5为JSON归一化比较方法流程图','图6为自动生成资源过滤规则判断流程图']:
    doc.add_paragraph(d, style='List Bullet')
doc.add_paragraph()

# INSERT FIGURES HERE - in Section 4
figs = ['fig1.png','fig2.png','fig3.png','fig4.png','fig5.png','fig6.png']
captions = ['图1 双向同步系统整体架构图','图2 正向同步方法流程图','图3 反向同步方法流程图',
            '图4 三级运行时字段智能清理方法','图5 JSON归一化比较方法','图6 自动生成资源过滤规则']
for fname, cap in zip(figs, captions):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(OUT_DIR, fname), width=Cm(14))
    c = doc.add_paragraph(cap)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].bold = True
    doc.add_paragraph()

# ===== Section 5 =====
doc.add_heading('五、具体实施方式', level=1)
doc.add_paragraph('实施例一：正向同步的变更预览与审核')
doc.add_paragraph('步骤S1，从GitLab仓库递归获取所有YAML文件。使用ListTree API配合Recursive参数，支持任意深度目录。')
doc.add_paragraph('步骤S2，使用sigs.k8s.io/yaml解析，确保数值为JSON兼容的float64类型，避免gopkg.in/yaml.v3产生的int类型导致DeepCopy panic。')
doc.add_paragraph('步骤S3，通过GVR解析器（内置映射表+Discovery API双层架构）确定资源的GroupVersionResource。')
doc.add_paragraph('步骤S4，对两侧资源执行三级清理。关键代码——数值类型感知匹配：')
doc.add_paragraph('  dropIfEqualInt(m, key, defaultVal) 通过type switch同时匹配int/int64/float64三种类型', style='List Bullet')
doc.add_paragraph('步骤S5，从K8s集群Get同名资源，同样执行三级清理。')
doc.add_paragraph('步骤S6，JSON归一化比较：递归将数值统一为float64，过滤空值，序列化后字符串比较。')
doc.add_paragraph('步骤S7，生成变更预览，两侧均为清理后的YAML，确保diff仅含真实业务变更。')
doc.add_paragraph('步骤S8，用户审核通过后，使用原始GitLab YAML通过Server-Side Apply应用到集群。')
doc.add_paragraph()
doc.add_paragraph('实施例二：反向同步的资源过滤与批量提交')
doc.add_paragraph('步骤S1，确定资源类型和命名空间。集群作用域资源全局列出一次，命名空间资源按ns分别列出。')
doc.add_paragraph('步骤S2，自动生成资源过滤：OwnerReferences检查→类型黑名单→系统资源白名单。')
doc.add_paragraph('步骤S3，三级清理后序列化为精简YAML，不含集群特定信息。')
doc.add_paragraph('步骤S4，与GitLab已有文件比较，相同则跳过。')
doc.add_paragraph('步骤S5，通过GitLab多文件原子提交API一次性提交所有变更。')
doc.add_paragraph()
doc.add_paragraph('实施例三：数值类型不一致问题的解决')
doc.add_paragraph('问题：K8s API返回int64(80)，YAML解析产生float64(80)，reflect.DeepEqual判定不等。')
doc.add_paragraph('方案：normalize函数递归将int/int32/int64/float32统一转为float64，JSON序列化后比较字符串。')

# ===== Section 6 =====
doc.add_heading('六、替代方案', level=1)
for a in ['1. YAML解析器：可用encoding/json先转JSON再解析；',
          '2. 比较算法：可实现递归类型无关DeepEqual；',
          '3. 清理策略：可通过OpenAPI Schema动态获取默认值；',
          '4. 同步协议：可用三方合并替代Server-Side Apply；',
          '5. 过滤规则：可基于标签选择器替代OwnerReferences。']:
    doc.add_paragraph(a)

# ===== Section 7 =====
doc.add_heading('七、其他事项', level=1)
doc.add_paragraph('C1、本申请所属项目公开时间：尚未公开')
doc.add_paragraph('C2、本申请期望完成初稿时间：2026年6月')

doc.save(OUTPUT)
print(f"Done! {OUTPUT}")
